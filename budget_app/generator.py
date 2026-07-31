from __future__ import annotations

import csv
import os
import shutil
import subprocess
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .storage import Item, Orcamento, iter_export_rows


TEMPLATE_DIR = Path("modelos")
OUTPUT_DIR = Path("saida")
EXPORT_PATH = Path("data/orcamentos_export.csv")


MONTHS_PT_BR = [
    "Janeiro",
    "Fevereiro",
    "Março",
    "Abril",
    "Maio",
    "Junho",
    "Julho",
    "Agosto",
    "Setembro",
    "Outubro",
    "Novembro",
    "Dezembro",
]


def parse_decimal(value: str) -> Decimal:
    cleaned = (value or "").strip().replace(".", "").replace(",", ".")
    if not cleaned:
        return Decimal("0")
    try:
        return Decimal(cleaned)
    except InvalidOperation as exc:
        raise ValueError(f"Valor numérico inválido: {value}") from exc


def only_digits(value: str) -> str:
    return "".join(char for char in (value or "") if char.isdigit())


def format_cpf(value: str) -> str:
    digits = only_digits(value)
    if not digits:
        return ""
    if len(digits) != 11:
        return value
    return f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"


def format_cell_phone(value: str) -> str:
    digits = only_digits(value)
    if not digits:
        return ""
    if len(digits) != 11:
        return value
    return f"({digits[:2]}) {digits[2:7]}-{digits[7:]}"


def require_valid_cpf(value: str) -> str:
    digits = only_digits(value)
    if not digits:
        return ""
    if len(digits) != 11:
        raise ValueError("CPF deve ter 11 dígitos.")
    return format_cpf(digits)


def require_valid_cell_phone(value: str) -> str:
    digits = only_digits(value)
    if not digits:
        return ""
    if len(digits) != 11:
        raise ValueError("Telefone celular deve ter 11 dígitos, incluindo DDD.")
    return format_cell_phone(digits)


def money(value: Decimal | str) -> str:
    decimal_value = parse_decimal(value) if isinstance(value, str) else value
    quantized = decimal_value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{quantized:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def quantity(value: Decimal) -> str:
    if value == value.to_integral():
        return str(int(value))
    return str(value).replace(".", ",")


def today_pt_br() -> str:
    now = date.today()
    return f"{now.day:02d} de {MONTHS_PT_BR[now.month - 1]} {now.year}"


def today_iso() -> str:
    return date.today().isoformat()


def date_for_document(value: str) -> str:
    value = (value or "").strip()
    if not value:
        return today_pt_br()
    try:
        parsed = datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return value
    return f"{parsed.day:02d} de {MONTHS_PT_BR[parsed.month - 1]} {parsed.year}"


def date_for_display(value: str) -> str:
    value = (value or "").strip()
    try:
        parsed = datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return value
    return parsed.strftime("%d/%m/%Y")


def safe_filename(text: str) -> str:
    allowed = []
    for char in text.strip().lower():
        if char.isalnum():
            allowed.append(char)
        elif char in (" ", "-", "_"):
            allowed.append("_")
    name = "".join(allowed).strip("_")
    while "__" in name:
        name = name.replace("__", "_")
    return name or "orcamento"


def available_templates() -> list[Path]:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    return sorted(
        path
        for path in TEMPLATE_DIR.glob("*.docx")
        if not path.name.startswith("~$")
    )


def get_template_path() -> Path:
    templates = available_templates()
    if not templates:
        raise FileNotFoundError(
            f"Nenhum modelo .docx encontrado em {TEMPLATE_DIR}. "
            "Coloque o novo modelo nessa pasta e valide antes de gerar."
        )
    return templates[0]


def ensure_template() -> Path:
    return get_template_path()


def validate_template() -> tuple[bool, list[str]]:
    messages: list[str] = []
    try:
        template_path = ensure_template()
        doc = Document(template_path)
        messages.append(f"OK: modelo ativo: {template_path}")
    except Exception as exc:
        return False, [f"Não foi possível abrir o template: {exc}"]

    paragraph_texts = [paragraph.text.strip() for paragraph in doc.paragraphs]
    joined = "\n".join(paragraph_texts)

    if "Dados Cliente" in joined:
        messages.append("OK: seção Dados Cliente encontrada.")
    else:
        messages.append("ERRO: seção Dados Cliente não encontrada.")

    try:
        for prefix in ("Cliente:", "CPF:", "Telefone:", "E-mail:"):
            replace_prefixed_paragraph(doc, prefix, "teste", after_text="Dados Cliente")
            messages.append(f"OK: campo {prefix} encontrado após Dados Cliente.")
    except Exception as exc:
        messages.append(f"ERRO: {exc}")

    if any(text.startswith("Campo Grande,") for text in paragraph_texts):
        messages.append("OK: campo de data encontrado.")
    else:
        messages.append('ERRO: campo de data não encontrado. Mantenha uma linha começando com "Campo Grande,".')

    if not doc.tables:
        messages.append("ERRO: tabela de produtos não encontrada.")
    else:
        table = doc.tables[0]
        if len(table.rows) >= 3 and len(table.columns) >= 7:
            messages.append("OK: tabela de produtos encontrada com estrutura suficiente.")
        else:
            messages.append("ERRO: a primeira tabela precisa ter pelo menos 3 linhas e 7 colunas.")

    ok = not any(message.startswith("ERRO:") for message in messages)
    return ok, messages


def set_runs_text(paragraph, texts: list[str]) -> None:
    while len(paragraph.runs) < len(texts):
        paragraph.add_run("")
    for idx, run in enumerate(paragraph.runs):
        run.text = texts[idx] if idx < len(texts) else ""


def set_paragraph_text(paragraph, text: str) -> None:
    set_runs_text(paragraph, [text])


def set_prefixed_paragraph_value(paragraph, prefix: str, value: str) -> None:
    value_text = f" {value}" if value else ""
    if len(paragraph.runs) >= 2:
        set_runs_text(paragraph, [prefix, value_text])
    else:
        set_runs_text(paragraph, [f"{prefix}{value_text}"])


def replace_prefixed_paragraph(
    doc: Document,
    prefix: str,
    value: str,
    after_text: str | None = None,
) -> None:
    active = after_text is None
    for paragraph in doc.paragraphs:
        if after_text and after_text in paragraph.text:
            active = True
            continue
        if not active:
            continue
        if paragraph.text.strip().startswith(prefix):
            set_prefixed_paragraph_value(paragraph, prefix, value)
            return
    raise ValueError(f"Campo não encontrado no template: {prefix}")


def replace_date_paragraph(doc: Document, localidade: str, data_orcamento: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Campo Grande,"):
            date_parts = data_orcamento.split(" ", 1)
            if len(paragraph.runs) >= 3 and len(date_parts) == 2:
                set_runs_text(paragraph, [f"{localidade}, ", f"{date_parts[0]} ", date_parts[1]])
            else:
                set_paragraph_text(paragraph, f"{localidade}, {data_orcamento}")
            return
    raise ValueError("Campo de data não encontrado no template.")


def insert_farmaceutico_responsavel(doc: Document, value: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Campo Grande,"):
            new_paragraph = paragraph.insert_paragraph_before()
            if paragraph._p.pPr is not None:
                new_paragraph._p.insert(0, deepcopy(paragraph._p.pPr))
            set_paragraph_text(new_paragraph, f"Farm. Resp.: {value}")
            if new_paragraph.runs:
                new_paragraph.runs[0].font.size = Pt(10)
            return
    raise ValueError("Campo de data não encontrado para inserir o farmacêutico responsável.")


def first_cell_run_properties(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run._r.rPr is not None:
                return deepcopy(run._r.rPr)
    return None


def apply_run_properties(run, run_properties) -> None:
    if run_properties is None:
        return
    current = run._r.rPr
    if current is not None:
        run._r.remove(current)
    run._r.insert(0, deepcopy(run_properties))


def fill_cell(cell, text: str, run_properties=None) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph_text(cell.paragraphs[0], text)
    if cell.paragraphs[0].runs:
        apply_run_properties(cell.paragraphs[0].runs[0], run_properties)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")


def format_run_properties(r_pr, *, bold: bool | None = None, size_pt: int | None = None) -> None:
    if bold is not None:
        for child in list(r_pr.findall(qn("w:b"))):
            r_pr.remove(child)
        if bold:
            bold_element = OxmlElement("w:b")
            r_pr.append(bold_element)
    if size_pt is not None:
        for tag in ("w:sz", "w:szCs"):
            for child in list(r_pr.findall(qn(tag))):
                r_pr.remove(child)
            size = OxmlElement(tag)
            size.set(qn("w:val"), str(size_pt * 2))
            r_pr.append(size)


def get_or_add_paragraph_run_properties(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    return r_pr


def format_cell_runs(cell, *, bold: bool | None = None, size_pt: int | None = None) -> None:
    for paragraph in cell.paragraphs:
        paragraph_r_pr = get_or_add_paragraph_run_properties(paragraph)
        format_run_properties(paragraph_r_pr, bold=bold, size_pt=size_pt)
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bold
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            format_run_properties(run._r.get_or_add_rPr(), bold=bold, size_pt=size_pt)


def remove_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def fill_products_table(doc: Document, itens: list[Item]) -> None:
    if not doc.tables:
        raise ValueError("Tabela de produtos não encontrada no template.")

    table = doc.tables[0]
    if len(table.rows) < 3:
        raise ValueError("A tabela de produtos precisa ter cabeçalho, item-modelo e total.")

    model_cells = table.rows[1].cells
    raw_column_run_properties = [first_cell_run_properties(cell) for cell in model_cells]
    fallback_run_properties = next(
        (properties for properties in raw_column_run_properties if properties is not None),
        None,
    )
    column_run_properties = [
        properties if properties is not None else fallback_run_properties
        for properties in raw_column_run_properties
    ]
    model_row_xml = deepcopy(table.rows[1]._tr)

    while len(table.rows) > 2:
        remove_row(table, 1)

    total_row = table.rows[-1]
    total_row._tr.addprevious(deepcopy(model_row_xml))
    remove_row(table, 1)

    for item in itens:
        new_row_xml = deepcopy(model_row_xml)
        total_row._tr.addprevious(new_row_xml)
        row = table.rows[-2]
        values = [
            item.codigo,
            item.produto,
            item.pmc,
            item.desconto,
            money(item.valor_unitario),
            quantity(item.quantidade),
            money(item.total),
        ]
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            fill_cell(cell, value, column_run_properties[idx])

    cells = total_row.cells
    seen = set()
    for cell in cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        fill_cell(cell, "")
    fill_cell(cells[0], "Total:", column_run_properties[0])
    fill_cell(cells[-1], money(sum((item.total for item in itens), Decimal("0"))), column_run_properties[-1])
    format_cell_runs(cells[0], bold=True, size_pt=10)
    format_cell_runs(cells[-1], bold=True, size_pt=10)


def render_orcamento(orcamento: Orcamento, sequence: int | None = None) -> Path:
    template_path = ensure_template()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(template_path)
    replace_prefixed_paragraph(doc, "Cliente:", orcamento.cliente_nome, after_text="Dados Cliente")
    replace_prefixed_paragraph(doc, "CPF:", format_cpf(orcamento.cpf), after_text="Dados Cliente")
    replace_prefixed_paragraph(doc, "Telefone:", format_cell_phone(orcamento.telefone), after_text="Dados Cliente")
    replace_prefixed_paragraph(doc, "E-mail:", orcamento.email or "", after_text="Dados Cliente")
    if orcamento.farmaceutico_responsavel:
        insert_farmaceutico_responsavel(doc, orcamento.farmaceutico_responsavel)
    replace_date_paragraph(doc, orcamento.localidade, date_for_document(orcamento.data_orcamento))
    fill_products_table(doc, orcamento.itens)

    seq = f"{sequence:04d}_" if sequence else ""
    filename = f"{seq}orcamento_{safe_filename(orcamento.cliente_nome)}.docx"
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def powershell_quote(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def convert_to_pdf_with_word(docx_path: Path) -> tuple[Path | None, str]:
    if os.name != "nt":
        return None, "Microsoft Word disponível apenas no Windows."

    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if not powershell:
        return None, "PowerShell não encontrado para acionar o Microsoft Word."

    source_path = docx_path.resolve()
    pdf_path = docx_path.with_suffix(".pdf")
    target_path = pdf_path.resolve()
    temp_path = target_path.with_name(f"{target_path.stem}.tmp.pdf")
    temp_path.unlink(missing_ok=True)

    script = f"""
$ErrorActionPreference = 'Stop'
$docxPath = {powershell_quote(str(source_path))}
$pdfPath = {powershell_quote(str(temp_path))}
if (-not (Test-Path -LiteralPath $docxPath)) {{
  throw "DOCX não encontrado: $docxPath"
}}
if (Test-Path -LiteralPath $pdfPath) {{
  Remove-Item -LiteralPath $pdfPath -Force
}}
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open($docxPath, $false, $true)
  $doc.SaveAs([ref] $pdfPath, [ref] 17)
}} finally {{
  if ($doc -ne $null) {{
    $doc.Close([ref] $false) | Out-Null
  }}
  if ($word -ne $null) {{
    $word.Quit() | Out-Null
  }}
  [System.GC]::Collect()
  [System.GC]::WaitForPendingFinalizers()
}}
"""
    result = subprocess.run(
        [
            powershell,
            "-NoProfile",
            "-NonInteractive",
            "-ExecutionPolicy",
            "Bypass",
            "-Command",
            script,
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode == 0 and temp_path.exists():
        temp_path.replace(target_path)
        return pdf_path, "PDF gerado com sucesso pelo Microsoft Word."

    temp_path.unlink(missing_ok=True)
    detail = (result.stderr or result.stdout or "erro desconhecido").strip()
    return None, f"Microsoft Word não gerou o PDF: {detail}"


def convert_to_pdf_with_libreoffice(docx_path: Path) -> tuple[Path | None, str]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, "LibreOffice/soffice não encontrado."

    profile_dir = Path("tmp/libreoffice-profile").resolve()
    profile_dir.mkdir(parents=True, exist_ok=True)
    source_path = docx_path.resolve()
    output_dir = docx_path.parent.resolve()
    pdf_path = docx_path.with_suffix(".pdf")
    temp_pdf_path = docx_path.with_suffix(".tmp.pdf")
    pdf_path.unlink(missing_ok=True)
    temp_pdf_path.unlink(missing_ok=True)

    try:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                f"-env:UserInstallation=file:///{profile_dir.as_posix()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(source_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except Exception as exc:
        return None, f"LibreOffice não gerou o PDF: {exc}"

    if result.returncode == 0 and pdf_path.exists():
        return pdf_path, "PDF gerado com sucesso pelo LibreOffice."
    detail = (result.stderr or result.stdout or "erro desconhecido").strip()
    return None, f"LibreOffice não gerou o PDF: {detail}"


def convert_to_pdf(docx_path: Path) -> tuple[Path | None, str]:
    attempts = [convert_to_pdf_with_word, convert_to_pdf_with_libreoffice]
    errors = []
    for converter in attempts:
        pdf_path, status = converter(docx_path)
        if pdf_path is not None:
            return pdf_path, status
        errors.append(status)
    return None, "PDF não gerado: " + " | ".join(errors)


def export_history_csv() -> Path:
    EXPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    rows = list(iter_export_rows())
    with EXPORT_PATH.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(
            [
                "orcamento_id",
                "criado_em",
                "cliente_nome",
                "farmaceutico_responsavel",
                "cpf",
                "telefone",
                "email",
                "data_orcamento",
                "localidade",
                "total_orcamento",
                "codigo",
                "produto",
                "pmc",
                "desconto",
                "valor_unitario",
                "quantidade",
                "total_item",
            ]
        )
        for row in rows:
            writer.writerow([row[key] for key in row.keys()])
    return EXPORT_PATH
